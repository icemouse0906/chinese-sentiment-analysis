import os
import pandas as pd
import matplotlib.pyplot as plt
import re
import unicodedata

def clean_text(s):
    if pd.isna(s):
        return ''
    s = str(s)
    s = re.sub(r'[\x00-\x1f\x7f-\x9f]', '', s)
    s = unicodedata.normalize('NFKC', s)
    return s
from docx import Document
from docx.shared import Inches

DATASETS = {
    '电商': 'NLP数据集/电商评论数据/online_shopping_10_cats.csv',
    '酒店': 'NLP数据集/酒店评论数据/ChnSentiCorp_htl_all.csv',
    '外卖': 'NLP数据集/外卖评论数据/waimai_10k.csv'
}
REPORT_PATH = 'experiment_report_final.docx'

# 设置macOS中文字体
plt.rcParams['font.sans-serif'] = ['PingFang SC', 'STHeiti', 'Arial Unicode MS', 'sans-serif']
plt.rcParams['axes.unicode_minus'] = False

def load_data(path, dataset_name=None):
    encodings = ['utf-8', 'gbk', 'gb18030']
    for enc in encodings:
        try:
            df = pd.read_csv(path, encoding=enc)
            break
        except UnicodeDecodeError:
            continue
    else:
        raise RuntimeError(f'无法读取 {path}，请检查文件编码')

    cols_lower = [c.lower() for c in df.columns]
    label_name_candidates = ['label', 'labels', 'sentiment', 'sentiment_label', 'polarity', 'rating', '评分', '类别', '情感']
    label_col = None
    for cand in label_name_candidates:
        for i, c in enumerate(cols_lower):
            if cand in c:
                label_col = df.columns[i]
                break
        if label_col is not None:
            break

    if label_col is None:
        n = len(df)
        unique_counts = {c: df[c].nunique() for c in df.columns}
        candidates = [c for c, u in unique_counts.items() if (u / max(1, n)) < 0.2 and u <= 50]
        if candidates:
            label_col = min(candidates, key=lambda c: unique_counts[c])

    if label_col is None:
        labels_found = None
        label_map = {
            '电商': 'ecommerce',
            '酒店': 'hotel',
            '外卖': 'waimai'
        }
        preferred_name = None
        if dataset_name and dataset_name in label_map:
            preferred_name = f'labels_{label_map[dataset_name]}.csv'
        if os.path.exists('output'):
            if preferred_name and preferred_name in os.listdir('output'):
                labels_found = preferred_name
            else:
                for f in os.listdir('output'):
                    if f.startswith('labels') and f.endswith('.csv'):
                        labels_found = f
                        break
        if labels_found:
            labels_df = pd.read_csv(os.path.join('output', labels_found), encoding='utf-8')
            if len(labels_df) == len(df):
                lbl_cols = list(labels_df.columns)
                lbl_cols_lower = [c.lower() for c in lbl_cols]
                label_name_candidates = ['label', 'labels', 'sentiment', 'sentiment_label', 'polarity', 'rating', '评分', '类别', '情感']
                chosen_lbl_col = None
                for cand in label_name_candidates:
                    for i, c in enumerate(lbl_cols_lower):
                        if cand in c:
                            chosen_lbl_col = lbl_cols[i]
                            break
                    if chosen_lbl_col:
                        break
                if chosen_lbl_col is None:
                    unique_counts_lbl = {c: labels_df[c].nunique() for c in lbl_cols}
                    small_candidates = [c for c,u in unique_counts_lbl.items() if u <= 50]
                    if small_candidates:
                        chosen_lbl_col = min(small_candidates, key=lambda c: unique_counts_lbl[c])
                if chosen_lbl_col is None and lbl_cols:
                    chosen_lbl_col = lbl_cols[-1]
                labels_series = labels_df[chosen_lbl_col].astype(str).reset_index(drop=True)
                possible_texts = [c for c in df.columns if df[c].dtype == object]
                text_col_guess = possible_texts[0] if possible_texts else df.columns[0]
                overlap_prop = df[text_col_guess].astype(str).isin(labels_series).mean()
                if overlap_prop > 0.5:
                    # 放弃自动匹配，返回 labels 为 None 以便上层走无标签分支
                    labels_series = None
                else:
                    df = df.reset_index(drop=True)
                    df['__label__'] = labels_series
                    label_col = '__label__'
            else:
                label_col = None
        # 如果仍未识别到标签列，不抛错；上层主流程会以 labels=None 处理无标签情形
        if label_col is None:
            return df, None

    df = df.dropna(subset=[label_col])
    text_col_candidates = [c for c in df.columns if any(k in c.lower() for k in ['text', 'review', 'content', '评论', '内容'])]
    text_col = text_col_candidates[0] if text_col_candidates else df.columns[0]
    return df[text_col].astype(str), df[label_col]

def analyze_and_plot_distribution(y, dataset_name):
    value_counts = pd.Series(y).value_counts().sort_values(ascending=False)
    value_counts.to_csv(f'results/{dataset_name}_label_distribution.csv', encoding='utf-8')
    plt.figure(figsize=(6,4))
    # 清洗标签，去除控制字符和NULL等非XML兼容字符
    def clean_text(s):
        if pd.isna(s):
            return ''
        s = str(s)
        # remove NULL and control chars
        s = re.sub(r'[\x00-\x1f\x7f-\x9f]', '', s)
        # normalize
        s = unicodedata.normalize('NFKC', s)
        return s

    labels = [clean_text(l)[:10] + ('...' if len(clean_text(l)) > 10 else '') for l in value_counts.index]
    plt.bar(labels, value_counts.values, color='skyblue')
    plt.title(f'{dataset_name} 类别分布')
    plt.xlabel('类别')
    plt.ylabel('样本数')
    plt.tight_layout()
    plt.savefig(f'results/{dataset_name}_label_distribution.png')
    plt.close()
    return value_counts

def append_eda_to_report(eda_infos):
    if not os.path.exists(REPORT_PATH):
        doc = Document()
    else:
        doc = Document(REPORT_PATH)
    doc.add_heading('数据探索（EDA）与类别分布分析', level=1)
    for dataset_name, value_counts in eda_infos:
        doc.add_heading(f'{dataset_name} 类别分布', level=2)
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Light List'
        table.cell(0,0).text = '类别'
        table.cell(0,1).text = '样本数'
        for idx, cnt in value_counts.items():
            row = table.add_row().cells
            clean_idx = clean_text(idx)
            row[0].text = clean_idx
            row[1].text = str(cnt)
        img_path = f'results/{dataset_name}_label_distribution.png'
        if os.path.exists(img_path):
            doc.add_picture(img_path, width=Inches(3.5))
        total = value_counts.sum()
        max_label, max_cnt = value_counts.index[0], value_counts.iloc[0]
        min_label, min_cnt = value_counts.index[-1], value_counts.iloc[-1]
        ratio = max_cnt / (min_cnt+1e-6)
        if ratio > 10:
            doc.add_paragraph(f'可以看出，{dataset_name} 存在严重的类别不平衡现象。例如，类别“{max_label}”有{max_cnt}个样本，而类别“{min_label}”仅有{min_cnt}个，比例高达{ratio:.1f}:1。')
        else:
            doc.add_paragraph(f'{dataset_name} 各类别分布较为均衡。')
    doc.save(REPORT_PATH)

def main():
    os.makedirs('results', exist_ok=True)
    eda_infos = []
    for dataset_name, path in DATASETS.items():
        if not os.path.exists(path):
            print(f'数据集 {dataset_name} 未找到，跳过')
            continue
        texts, labels = load_data(path, dataset_name=dataset_name)
        if labels is None:
            # 没有标签，输出文本长度分布与示例
            lens = texts.map(len)
            lens.to_csv(f'results/{dataset_name}_text_length.csv', encoding='utf-8')
            plt.figure(figsize=(6,4))
            lens.hist(bins=50)
            plt.title(f'{dataset_name} 文本长度分布')
            plt.savefig(f'results/{dataset_name}_text_length.png')
            plt.close()
            eda_infos.append((dataset_name, pd.Series({'no_label': len(texts)})))
        else:
            value_counts = analyze_and_plot_distribution(labels, dataset_name)
            eda_infos.append((dataset_name, value_counts))
    append_eda_to_report(eda_infos)
    print('EDA分析与报告追加完成！')

if __name__ == '__main__':
    main()
