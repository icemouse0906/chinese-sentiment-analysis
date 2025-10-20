# coding: utf-8
"""
Append error analysis and feature importance to the final docx report
"""
from pathlib import Path
import pandas as pd
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn

ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / 'output'
FINAL = OUT / 'experiment_report_final.docx'
TMP = OUT / 'experiment_report_final_tmp.docx'

# load existing final docx (if exists)
if not FINAL.exists():
    print('Missing', FINAL)
    raise SystemExit(1)

doc = Document(str(FINAL))
# add a section heading
doc.add_page_break()
doc.add_heading('错误分析与特征重要性（追加）', level=1)

DATASETS = ['hotel','ecommerce','waimai']
for ds in DATASETS:
    doc.add_heading(f'数据集：{ds}', level=2)
    # feature importance SVM
    f_svm = OUT / f'feature_importance_{ds}_svm.csv'
    f_nb = OUT / f'feature_importance_{ds}_nb.csv'
    if f_svm.exists():
        df = pd.read_csv(f_svm)
        doc.add_paragraph('SVM: 正面/负面权重最高的词（coef）')
        # create small table: 10 rows, 4 cols
        table = doc.add_table(rows=1, cols=4)
        hdr = table.rows[0].cells
        hdr[0].text = 'pos_word'
        hdr[1].text = 'pos_coef'
        hdr[2].text = 'neg_word'
        hdr[3].text = 'neg_coef'
        for _, r in df.iterrows():
            row = table.add_row().cells
            row[0].text = str(r.get('pos_word',''))
            row[1].text = f"{r.get('pos_coef',''):.4f}"
            row[2].text = str(r.get('neg_word',''))
            row[3].text = f"{r.get('neg_coef',''):.4f}"
    if f_nb.exists():
        dfnb = pd.read_csv(f_nb)
        doc.add_paragraph('NB: 正面/负面对数概率高的词（feature_log_prob）')
        table = doc.add_table(rows=1, cols=4)
        hdr = table.rows[0].cells
        hdr[0].text = 'pos_word'
        hdr[1].text = 'pos_logprob'
        hdr[2].text = 'neg_word'
        hdr[3].text = 'neg_logprob'
        for _, r in dfnb.iterrows():
            row = table.add_row().cells
            row[0].text = str(r.get('pos_word',''))
            row[1].text = f"{r.get('pos_logprob',''):.4f}"
            row[2].text = str(r.get('neg_word',''))
            row[3].text = f"{r.get('neg_logprob',''):.4f}"

    # error analysis: show 4 representative misclassified samples from SVM
    e_svm = OUT / f'error_analysis_{ds}_svm.csv'
    e_nb = OUT / f'error_analysis_{ds}_nb.csv'
    doc.add_heading('误判样例（SVM）', level=3)
    if e_svm.exists():
        edf = pd.read_csv(e_svm)
        # sample up to 5
        sample = edf.sample(n=min(5, len(edf)), random_state=42)
        for _, r in sample.iterrows():
            doc.add_paragraph(f"index={r.get('index')} fold={r.get('fold')} true={r.get('true')} pred={r.get('pred')}")
            doc.add_paragraph(str(r.get('review','')))
            doc.add_paragraph('tokens_join: ' + str(r.get('tokens_join','')))
            doc.add_paragraph('---')
    else:
        doc.add_paragraph('No SVM error analysis file.')

    doc.add_heading('误判样例（NB）', level=3)
    if e_nb.exists():
        edf = pd.read_csv(e_nb)
        sample = edf.sample(n=min(5, len(edf)), random_state=42)
        for _, r in sample.iterrows():
            doc.add_paragraph(f"index={r.get('index')} fold={r.get('fold')} true={r.get('true')} pred={r.get('pred')}")
            doc.add_paragraph(str(r.get('review','')))
            doc.add_paragraph('tokens_join: ' + str(r.get('tokens_join','')))
            doc.add_paragraph('---')
    else:
        doc.add_paragraph('No NB error analysis file.')

# Save updated doc
doc.save(TMP)
# replace
import os
os.replace(str(TMP), str(FINAL))
print('Appended analysis and updated', FINAL)
