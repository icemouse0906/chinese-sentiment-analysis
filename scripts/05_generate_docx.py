# coding: utf-8
"""
Generate a Word (.docx) report from the Markdown experiment report and include key images and CV summary tables.
"""
from pathlib import Path
from docx import Document
from docx.shared import Inches
import pandas as pd

ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / 'output'
MD = OUT / 'experiment_report.md'
DOCX = OUT / 'experiment_report.docx'

# Helper: add a markdown-like header or paragraph
def add_paragraph_md(doc, line):
    line = line.rstrip('\n')
    if line.startswith('# '):
        doc.add_heading(line[2:].strip(), level=1)
    elif line.startswith('## '):
        doc.add_heading(line[3:].strip(), level=2)
    elif line.startswith('### '):
        doc.add_heading(line[4:].strip(), level=3)
    elif line.strip() == '':
        doc.add_paragraph('')
    else:
        doc.add_paragraph(line)

# Read markdown and create document
doc = Document()
if not MD.exists():
    print('Missing', MD)
    raise SystemExit(1)

with MD.open('r', encoding='utf-8') as f:
    for line in f:
        add_paragraph_md(doc, line)

# Insert some key images if available
# length hist images
for img in (OUT / 'length_hist_hotel.png', OUT / 'length_hist_ecommerce.png', OUT / 'length_hist_waimai.png'):
    if img.exists():
        doc.add_page_break()
        doc.add_heading(f'Length histogram: {img.name}', level=2)
        doc.add_picture(str(img), width=Inches(6))

# Insert CV summary tables and all fold images with captions
cv_root = OUT / 'cv'
if cv_root.exists():
    for ds in ['hotel','ecommerce','waimai']:
        ms = cv_root / ds / 'metrics_summary.csv'
        doc.add_page_break()
        doc.add_heading(f'Cross-validation results: {ds}', level=1)
        if ms.exists():
            df = pd.read_csv(ms)
            doc.add_paragraph('Summary (mean ± std) for precision/recall/f1/accuracy:')
            # format table with 1 header + rows
            table = doc.add_table(rows=1, cols=len(df.columns))
            hdr_cells = table.rows[0].cells
            for i, c in enumerate(df.columns):
                hdr_cells[i].text = c
            for _, r in df.iterrows():
                row_cells = table.add_row().cells
                for i, c in enumerate(df.columns):
                    row_cells[i].text = str(r[c])
        else:
            doc.add_paragraph('No metrics_summary.csv found for ' + ds)

        # Insert per-fold images
        ds_cv = cv_root / ds
        if ds_cv.exists():
            # try to include metrics_summary to craft small caption
            summary_text = ''
            try:
                summary_df = pd.read_csv(ds_cv / 'metrics_summary.csv')
                summary_text = '\n'.join([f"{row['model']}: f1_mean={row['f1_mean']:.3f} ± {row['f1_std']:.3f}" for _, row in summary_df.iterrows()])
                if summary_text:
                    doc.add_paragraph('Summary:')
                    doc.add_paragraph(summary_text)
            except Exception:
                pass

            # iterate folds
            # detect max fold number by scanning filenames
            imgs = list(ds_cv.glob('fold_*_svm_confusion.png'))
            max_folds = max([int(p.name.split('_')[1]) for p in imgs]) if imgs else 5
            for i in range(1, max_folds+1):
                doc.add_heading(f'Fold {i}', level=2)
                # insert SVM confusion
                cm_path = ds_cv / f'fold_{i}_svm_confusion.png'
                if cm_path.exists():
                    doc.add_paragraph('SVM confusion matrix (rows=true labels, cols=predicted labels):')
                    doc.add_picture(str(cm_path), width=Inches(5))
                # insert PR/ROC for SVM and NB
                for model in ['svm','nb']:
                    pr = ds_cv / f'fold_{i}_{model}_pr.png'
                    roc = ds_cv / f'fold_{i}_{model}_roc.png'
                    if pr.exists():
                        doc.add_paragraph(f'{model.upper()} PR curve (fold {i}):')
                        doc.add_picture(str(pr), width=Inches(5))
                    if roc.exists():
                        doc.add_paragraph(f'{model.upper()} ROC curve (fold {i}):')
                        doc.add_picture(str(roc), width=Inches(5))
        else:
            doc.add_paragraph('No CV folder found for ' + ds)

# Insert a reference to samples_for_annotation files
for samp in ['samples_for_annotation_hotel.csv','samples_for_annotation_ecommerce.csv','samples_for_annotation_waimai.csv']:
    p = OUT / samp
    if p.exists():
        doc.add_page_break()
        doc.add_heading(f'Sample for annotation: {p.name}', level=2)
        # include first 10 lines as example
        with p.open('r', encoding='utf-8') as f:
            for i, line in enumerate(f):
                if i > 10:
                    break
                doc.add_paragraph(line.strip())

# Save
doc.save(DOCX)
print('Wrote', DOCX)
