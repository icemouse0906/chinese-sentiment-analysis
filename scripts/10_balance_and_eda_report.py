import os
import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
import re
import unicodedata
from sklearn.model_selection import GridSearchCV, StratifiedKFold, train_test_split
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.naive_bayes import MultinomialNB
from sklearn.svm import LinearSVC
from sklearn.metrics import classification_report, confusion_matrix, roc_curve, precision_recall_curve, auc
from imblearn.over_sampling import SMOTE, RandomOverSampler
from docx import Document
from docx.shared import Inches
import warnings
import time
from datetime import datetime

# --- 进度日志 ---
def log(message):
    print(f"[{datetime.now().strftime('%H:%M:%S')}] {message}")

warnings.filterwarnings('ignore')

# 数据集路径
DATASETS = {
    '电商': 'NLP数据集/电商评论数据/online_shopping_10_cats.csv',
    '酒店': 'NLP数据集/酒店评论数据/ChnSentiCorp_htl_all.csv',
    '外卖': 'NLP数据集/外卖评论数据/waimai_10k.csv'
}

MODELS = {
    'SVM': (LinearSVC(dual=True), {'C': [0.1, 1, 5, 10], 'max_iter': [1000, 3000]}),
    'NB': (MultinomialNB(), {'alpha': [0.1, 0.5, 1.0, 2.0]})
}

REPORT_PATH = 'experiment_report_final.docx'

# 自动编码检测与回退
def load_data(path, dataset_name=None):
    encodings = ['utf-8', 'gbk', 'gb18030']
    for enc in encodings:
        try:
            df = pd.read_csv(path, encoding=enc)
            break
        except UnicodeDecodeError:
            continue
    else:
        raise RuntimeError(f'无法读取 {path}，请检查文件编码')

    cols_lower = [c.lower() for c in df.columns]
    # 优先通过列名匹配标签列
    label_name_candidates = ['label', 'labels', 'sentiment', 'sentiment_label', 'polarity', 'rating', '评分', '类别', '情感']
    label_col = None
    for cand in label_name_candidates:
        for i, c in enumerate(cols_lower):
            if cand in c:
                label_col = df.columns[i]
                break
        if label_col is not None:
            break

    # 优先通过列名匹配文本列
    text_name_candidates = ['text', 'review', 'content', '评论', '内容']
    text_col = None
    for cand in text_name_candidates:
        for i, c in enumerate(cols_lower):
            if cand in c:
                text_col = df.columns[i]
                break
        if text_col is not None:
            break

    # 如果没有找到明确的标签列，使用唯一值比例进行推断
    if label_col is None:
        n = len(df)
        unique_counts = {c: df[c].nunique() for c in df.columns}
        # 选择唯一值远小于总样本数的列作为标签（例如 unique/n < 0.2 且 unique <= 50）
        candidates = [c for c, u in unique_counts.items() if (u / max(1, n)) < 0.2 and u <= 50]
        if candidates:
            # 选择唯一值最小的作为标签列
            label_col = min(candidates, key=lambda c: unique_counts[c])

    # 如果没有找到文本列，选择字符串平均长度较大的列作为文本列（避免选成标签）
    if text_col is None:
        obj_cols = [c for c in df.columns if df[c].dtype == object]
        if obj_cols:
            avg_len = {c: df[c].astype(str).map(len).mean() for c in obj_cols}
            text_col = max(avg_len, key=avg_len.get)

    # 最后检查标签列和文本列是否合理
    if label_col is None:
        # 尝试在 output/ 目录寻找已生成的 labels 文件，优先按数据集映射名匹配
        label_map = {
            '电商': 'ecommerce',
            '酒店': 'hotel',
            '外卖': 'waimai'
        }
        labels_found = None
        preferred_name = None
        if dataset_name and dataset_name in label_map:
            preferred_name = f'labels_{label_map[dataset_name]}.csv'
        if os.path.exists('output'):
            # 优先使用映射到的文件名
            if preferred_name and preferred_name in os.listdir('output'):
                labels_found = preferred_name
            else:
                # 否则选择第一个以 labels 开头的文件（保守选择）
                for f in os.listdir('output'):
                    if f.startswith('labels') and f.endswith('.csv'):
                        labels_found = f
                        break
        if labels_found:
            labels_df = pd.read_csv(os.path.join('output', labels_found), encoding='utf-8')
            # 仅在行数匹配时尝试对齐
            if len(labels_df) == len(df):
                # 尝试在 labels_df 中选择合适的标签列：优先列名匹配，其次最小唯一值
                lbl_cols = list(labels_df.columns)
                lbl_cols_lower = [c.lower() for c in lbl_cols]
                label_name_candidates = ['label', 'labels', 'sentiment', 'sentiment_label', 'polarity', 'rating', '评分', '类别', '情感']
                chosen_lbl_col = None
                for cand in label_name_candidates:
                    for i, c in enumerate(lbl_cols_lower):
                        if cand in c:
                            chosen_lbl_col = lbl_cols[i]
                            break
                    if chosen_lbl_col:
                        break
                if chosen_lbl_col is None:
                    # 选择唯一值最小且不超过 50 的那一列
                    unique_counts_lbl = {c: labels_df[c].nunique() for c in lbl_cols}
                    small_candidates = [c for c,u in unique_counts_lbl.items() if u <= 50]
                    if small_candidates:
                        chosen_lbl_col = min(small_candidates, key=lambda c: unique_counts_lbl[c])
                # 如果仍然没有找到合适的标签列，尝试最后一列作为退路
                if chosen_lbl_col is None and lbl_cols:
                    chosen_lbl_col = lbl_cols[-1]

                # 将 labels_series 与主 df 对齐前，进行简单检查以避免把文本当标签
                labels_series = labels_df[chosen_lbl_col].astype(str).reset_index(drop=True)
                # 估算 labels_series 与可能的文本列重合比例（如果高，说明 labels 文件可能包含文本）
                possible_texts = [c for c in df.columns if df[c].dtype == object]
                text_col_guess = possible_texts[0] if possible_texts else df.columns[0]
                overlap_prop = df[text_col_guess].astype(str).isin(labels_series).mean()
                if overlap_prop > 0.5:
                    # 高重合，说明 labels 列可能是文本，放弃自动匹配
                    labels_series = None
                else:
                    df = df.reset_index(drop=True)
                    df['__label__'] = labels_series
                    label_col = '__label__'
            else:
                # 行数不匹配，放弃自动匹配
                label_col = None

    if label_col is None:
        raise RuntimeError(f'未能自动识别标签列，请手动检查数据并在脚本中指定标签列。可用列: {list(df.columns)}')
    if text_col is None:
        # 退回到第一列作为文本
        text_col = df.columns[0]

    # 避免把文本列误认为标签列
    if label_col == text_col:
        # 如果标签列与文本列相同，说明识别失败，抛出错误
        raise RuntimeError(f'检测到标签列与文本列相同（{label_col}），请检查数据或手动指定列')

    df = df.dropna(subset=[label_col])
    return df[text_col].astype(str), df[label_col]

def clean_text(s):
    if pd.isna(s):
        return ''
    s = str(s)
    s = re.sub(r'[\x00-\x1f\x7f-\x9f]', '', s)
    s = unicodedata.normalize('NFKC', s)
    return s

# 快速模式：用于教学场景，避免长时间运行
QUICK_MODE = False

# 类别分布分析与可视化
def analyze_and_plot_distribution(y, dataset_name):
    import matplotlib.pyplot as plt
    plt.rcParams['font.sans-serif'] = ['PingFang SC', 'STHeiti', 'Arial Unicode MS', 'sans-serif']  # macOS中文字体回退
    plt.rcParams['axes.unicode_minus'] = False
    value_counts = pd.Series(y).value_counts().sort_values(ascending=False)
    value_counts.to_csv(f'results/{dataset_name}_label_distribution.csv', encoding='utf-8')
    plt.figure(figsize=(6,4))
    # 标签只显示前10个字符，避免长文本导致异常
    labels = [clean_text(l)[:10] + ('...' if len(clean_text(l)) > 10 else '') for l in value_counts.index]
    plt.bar(labels, value_counts.values, color='skyblue')
    plt.title(f'{dataset_name} 类别分布')
    plt.xlabel('类别')
    plt.ylabel('样本数')
    plt.tight_layout()
    plt.savefig(f'results/{dataset_name}_label_distribution.png')
    plt.close()
    return value_counts

# 过采样处理
def balance_data(X, y):
    if len(set(y)) < 2:
        return X, y, '类别数不足2，无法平衡'
    
    vc = pd.Series(y).value_counts()
    majority_class_label = vc.index[0]
    majority_count = vc.iloc[0]
    target_ratio = 0.5

    # 统一构造 sampling_strategy 为字典：对每个非多数类，目标数 = int(majority_count * target_ratio)
    strategy_dict = {}
    for label, count in vc.items():
        if label == majority_class_label:
            continue
        target_count = int(majority_count * target_ratio)
        if target_count > count:
            strategy_dict[label] = target_count

    # 如果字典为空（例如所有其他类已达到目标），则无需采样
    if not strategy_dict:
        return X, y, '所有少数类样本已达标，无需过采样'

    sampling_strategy = strategy_dict
    method = f'过采样(目标=多数类的{int(target_ratio*100)}%)'

    min_count = vc.min()
    # 如果 X 是文本（object / str），不要使用 SMOTE（会尝试把文本转为 float）；改用 RandomOverSampler（复制样本）
    is_text = False
    try:
        # pandas Series 检测
        is_text = pd.Series(X).dtype == object or (len(X) > 0 and isinstance(X[0], str))
    except Exception:
        is_text = True

    # 如果 X 是文本（object / str），不要使用 SMOTE（会尝试把文本转为 float）；改用 RandomOverSampler（复制样本）
    if is_text:
        sampler = RandomOverSampler(sampling_strategy=sampling_strategy, random_state=42)
        method = '随机过采样(文本复制) - ' + method
        X_array = np.array(X).reshape(-1, 1)
        X_res, y_res = sampler.fit_resample(X_array, y)
        return X_res.ravel(), y_res, method

    # 否则尝试使用 SMOTE（数值特征）作为优选，若不满足 SMOTE 的条件则退回 RandomOverSampler
    if min_count < 6:
        sampler = RandomOverSampler(sampling_strategy=sampling_strategy, random_state=42)
        method = '随机' + method
    else:
        try:
            sampler = SMOTE(sampling_strategy=sampling_strategy, random_state=42)
            method = 'SMOTE' + method
        except Exception:
            sampler = RandomOverSampler(sampling_strategy=sampling_strategy, random_state=42)
            method = '随机' + method

    X_res, y_res = sampler.fit_resample(np.array(X), y)
    return X_res, y_res, method

# 调优与评估
def tune_and_evaluate(X, y, model, param_grid, dataset_name, model_name, use_split=False):
    tfidf = TfidfVectorizer(max_features=5000)
    X_vec = tfidf.fit_transform(X)
    if use_split:
        # 留出法
        X_train, X_test, y_train, y_test = train_test_split(X_vec, y, test_size=0.3, random_state=42, stratify=y)
        model.fit(X_train, y_train)
        y_pred = model.predict(X_test)
        report = classification_report(y_test, y_pred, output_dict=True)
        report_df = pd.DataFrame(report).transpose()
        report_path = f'results/{dataset_name}_{model_name}_report.csv'
        os.makedirs('results', exist_ok=True)
        report_df.to_csv(report_path, encoding='utf-8')
        cm = confusion_matrix(y_test, y_pred)
    else:
        skf = StratifiedKFold(n_splits=5, shuffle=True, random_state=42)
        n_jobs = 1 if QUICK_MODE else -1
        grid = GridSearchCV(model, param_grid, cv=skf, scoring='f1_weighted', n_jobs=n_jobs)
        grid.fit(X_vec, y)
        best_model = grid.best_estimator_
        y_pred = best_model.predict(X_vec)
        report = classification_report(y, y_pred, output_dict=True)
        report_df = pd.DataFrame(report).transpose()
        report_path = f'results/{dataset_name}_{model_name}_report.csv'
        os.makedirs('results', exist_ok=True)
        report_df.to_csv(report_path, encoding='utf-8')
        cm = confusion_matrix(y, y_pred)
    # 混淆矩阵
    plt.figure(figsize=(5,4))
    plt.imshow(cm, cmap='Blues')
    plt.title(f'{dataset_name}-{model_name} 混淆矩阵')
    plt.xlabel('预测')
    plt.ylabel('真实')
    plt.colorbar()
    plt.savefig(f'results/{dataset_name}_{model_name}_cm.png', bbox_inches='tight')
    plt.close()
    return report_df

# 追加到Word报告
def append_to_report(eda_infos, results):
    if not os.path.exists(REPORT_PATH):
        doc = Document()
    else:
        doc = Document(REPORT_PATH)
    # EDA部分
    doc.add_heading('数据探索（EDA）与类别分布分析', level=1)
    for dataset_name, value_counts in eda_infos:
        doc.add_heading(f'{dataset_name} 类别分布', level=2)
        
        # 只在类别数合理时（< 100）才写入详细表格，避免把文本误当类别导致超大表格
        num_classes = len(value_counts)
        if num_classes < 100:
            table = doc.add_table(rows=1, cols=2)
            table.style = 'Light List'
            table.cell(0,0).text = '类别'
            table.cell(0,1).text = '样本数'
            for idx, cnt in value_counts.items():
                row = table.add_row().cells
                row[0].text = clean_text(str(idx))[:50]  # 限制长度避免超长文本
                row[1].text = str(cnt)
        else:
            # 类别过多（可能是误把文本当类别），只写汇总统计
            doc.add_paragraph(f'警告：检测到 {num_classes} 个唯一类别（可能为文本数据被误识为类别），仅显示汇总统计。')
            doc.add_paragraph(f'样本总数：{value_counts.sum()}')
            doc.add_paragraph(f'最多样本的类别：{value_counts.iloc[0]} 个样本')
            doc.add_paragraph(f'最少样本的类别：{value_counts.iloc[-1]} 个样本')
        
        img_path = f'results/{dataset_name}_label_distribution.png'
        if os.path.exists(img_path):
            doc.add_picture(img_path, width=Inches(3.5))
        # 分析性文字
        total = value_counts.sum()
        max_label, max_cnt = clean_text(value_counts.index[0]), value_counts.iloc[0]
        min_label, min_cnt = clean_text(value_counts.index[-1]), value_counts.iloc[-1]
        ratio = max_cnt / (min_cnt+1e-6)
        # 结构化写法：发现→分析→解决→结论
        doc.add_heading(f'{dataset_name}：发现与分析', level=3)
        if ratio > 10:
            doc.add_paragraph(f'发现：{dataset_name} 存在严重的类别不平衡，类别“{max_label}”占主导（{max_cnt} / {total}），而“{min_label}”仅有 {min_cnt} 条样本，比例约 {ratio:.1f}:1。')
            doc.add_paragraph('分析：这种不平衡将使模型倾向于多数类，导致少数类（如负面评论）被忽略，评价指标（如F1）对少数类不敏感。')
        else:
            doc.add_paragraph(f'发现：{dataset_name} 各类别分布相对均衡（样本总数 {total}）。')
            doc.add_paragraph('分析：可直接进行交叉验证与常规调参。')
    # 模型与采样部分
    doc.add_heading('数据不平衡处理与模型优化', level=1)
    for dataset_name, model_name, method, report_df, strategy in results:
        doc.add_heading(f'{dataset_name} - {model_name}', level=2)
        doc.add_paragraph(f'解决方案：针对类别不平衡，采用了 {method}。')
        doc.add_paragraph(f'验证策略：{strategy}')
        table = doc.add_table(rows=1, cols=len(report_df.columns)+1)
        table.style = 'Light List'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = '类别'
        for i, col in enumerate(report_df.columns):
            hdr_cells[i+1].text = col
        for idx, row in report_df.iterrows():
            cells = table.add_row().cells
            cells[0].text = clean_text(idx)
            for i, val in enumerate(row):
                cells[i+1].text = f'{val:.3f}' if isinstance(val, float) else clean_text(val)
        img_cm = f'results/{dataset_name}_{model_name}_cm.png'
        if os.path.exists(img_cm):
            doc.add_picture(img_cm, width=Inches(3.5))
    doc.save(REPORT_PATH)

# 主流程
def main():
    os.makedirs('results', exist_ok=True)
    eda_infos = []
    results = []
    
    total_start_time = time.time()
    log("实验开始...")

    for dataset_name, path in DATASETS.items():
        dataset_start_time = time.time()
        log(f"--- 开始处理【{dataset_name}】数据集 ---")

        if not os.path.exists(path):
            log(f"【{dataset_name}】数据集未找到，跳过")
            continue
        
        # --- 数据加载 ---
        load_start = time.time()
        X, y = load_data(path, dataset_name=dataset_name)
        log(f"【{dataset_name}】数据加载完成 (耗时: {time.time() - load_start:.2f}s)")

        # --- EDA ---
        eda_start = time.time()
        value_counts = analyze_and_plot_distribution(y, dataset_name)
        eda_infos.append((dataset_name, value_counts))
        log(f"【{dataset_name}】EDA与类别分布图生成完成 (耗时: {time.time() - eda_start:.2f}s)")

        # --- 过采样 ---
        balance_start = time.time()
        X_bal, y_bal, method = balance_data(X, y)
        log(f"【{dataset_name}】数据不平衡处理完成 (方法: {method}, 耗时: {time.time() - balance_start:.2f}s)")
        log(f"【{dataset_name}】采样后样本数: {len(y_bal)}")

        # --- 模型调优 ---
        min_count = pd.Series(y_bal).value_counts().min()
        for model_name, (model, param_grid) in MODELS.items():
            model_start_time = time.time()
            log(f"【{dataset_name}】开始 {model_name} 模型调优...")
            
            if min_count < 5:
                strategy = '留出法（Train-Test Split）'
                report_df = tune_and_evaluate(X_bal, y_bal, model, param_grid, dataset_name, model_name, use_split=True)
            else:
                strategy = '5折交叉验证（StratifiedKFold）'
                report_df = tune_and_evaluate(X_bal, y_bal, model, param_grid, dataset_name, model_name, use_split=False)
            
            results.append((dataset_name, model_name, method, report_df, strategy))
            log(f"【{dataset_name}】{model_name} 模型调优完成 (耗时: {time.time() - model_start_time:.2f}s)")
        
        log(f"--- 【{dataset_name}】数据集处理完成 (总耗时: {time.time() - dataset_start_time:.2f}s) ---\n")

    # --- 报告生成 ---
    report_start_time = time.time()
    log("开始生成最终报告...")
    append_to_report(eda_infos, results)
    log(f"报告生成完成 (耗时: {time.time() - report_start_time:.2f}s)")
    
    log(f"所有任务完成！(总耗时: {time.time() - total_start_time:.2f}s)")

if __name__ == '__main__':
    main()
