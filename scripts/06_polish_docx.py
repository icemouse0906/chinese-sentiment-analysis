# coding: utf-8
"""
Polish the Word report: create a cover page, contents placeholder, executive summary,
per-dataset CV sections with metrics table and images (with captions), and appendix.
Produces output/experiment_report_final.docx
"""
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
import pandas as pd
import datetime

ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / 'output'
MD = OUT / 'experiment_report.md'
CV_ROOT = OUT / 'cv'
FINAL = OUT / 'experiment_report_final.docx'

# Read markdown to extract abstract and main findings
def extract_section(md_lines, title):
    start = None
    for i, line in enumerate(md_lines):
        if line.strip().startswith(title):
            start = i+1
            break
    if start is None:
        return ''
    # collect until next top-level heading '##'
    collected = []
    for line in md_lines[start:]:
        if line.startswith('## '):
            break
        collected.append(line.rstrip('\n'))
    return '\n'.join(collected).strip()

# load md
if not MD.exists():
    print('Missing', MD)
    raise SystemExit(1)
md_lines = MD.read_text(encoding='utf-8').splitlines()
abstract = extract_section(md_lines, '## 一、摘要')
main_findings = extract_section(md_lines, '## 十、主要发现与讨论')

# create document
doc = Document()
# set default font to a common Chinese-friendly font if available
style = doc.styles['Normal']
font = style.font
font.name = 'Microsoft YaHei'
font.size = Pt(11)
# for Chinese font compatibility in Word, set _element rPr rFonts
r = style.element.rPr
r.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

# Cover page
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
run = p.add_run('实验报告（作业二）')
run.font.size = Pt(28)
run.bold = True
p = doc.add_paragraph()
p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
p.add_run('\n')

meta = doc.add_paragraph()
meta.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
meta.add_run(f'Date: {datetime.date.today().isoformat()}\n')
meta.add_run('Author: 自动生成（请校验）')

doc.add_page_break()

# Contents placeholder
doc.add_heading('目录 (Contents)', level=1)
doc.add_paragraph('说明：在 Microsoft Word 中可右键更新目录以获得页码。此处为自动生成的目录占位，若需精确页码请在 Word 中更新 TOC。')
# Simple contents (we will list main sections)
sections = ['摘要', '方法与预处理', '模型与评估', 'Cross-validation results (hotel)', 'Cross-validation results (ecommerce)', 'Cross-validation results (waimai)', '样本抽样 (annotation)', '附录：输出文件列表']
for s in sections:
    p = doc.add_paragraph()
    p.add_run(s)
    # leave space for page number
    p.add_run(' ............................................. ')

doc.add_page_break()

# Executive Summary
doc.add_heading('摘要 / Executive Summary', level=1)
if abstract:
    for para in abstract.split('\n'):
        doc.add_paragraph(para)
else:
    doc.add_paragraph('无摘要内容（请查看 experiment_report.md）')

doc.add_heading('主要发现与讨论', level=2)
if main_findings:
    for para in main_findings.split('\n'):
        doc.add_paragraph(para)
else:
    doc.add_paragraph('无主要发现节（请查看 experiment_report.md）')

# Add a short methodology summary
doc.add_heading('方法与预处理（简述）', level=2)
doc.add_paragraph('数据：酒店、商品、外卖三套中文评论数据。\n预处理：自动编码检测+回退，unicode_escape 解码，jieba 分词，TF-IDF 特征化。\n标签：使用 SnowNLP 生成 sentiment_score (0-1)，阈值 0.5 作为伪标签。\n模型：MultinomialNB 与 LinearSVC(class_weight=\'balanced\')。训练集上采样用于缓解不平衡。')

# Per-dataset CV sections with better captions
for ds in ['hotel','ecommerce','waimai']:
    doc.add_page_break()
    doc.add_heading(f'Cross-validation results — {ds}', level=1)
    ms = CV_ROOT / ds / 'metrics_summary.csv'
    if ms.exists():
        df = pd.read_csv(ms)
        doc.add_paragraph('Metrics summary (see metrics_per_fold.csv for per-fold details):')
        table = doc.add_table(rows=1, cols=len(df.columns))
        hdr = table.rows[0].cells
        for i,c in enumerate(df.columns):
            hdr[i].text = c
        for _, row in df.iterrows():
            r = table.add_row().cells
            for i,c in enumerate(df.columns):
                r[i].text = str(row[c])
    else:
        doc.add_paragraph('No metrics_summary.csv found for ' + ds)

    # insert folds with captions
    ds_cv = CV_ROOT / ds
    if ds_cv.exists():
        imgs = sorted(list(ds_cv.glob('fold_*_svm_confusion.png')))
        if not imgs:
            # fallback: process fold numbers 1..5
            folds = range(1,6)
        else:
            folds = [int(p.name.split('_')[1]) for p in imgs]
        for i in folds:
            doc.add_heading(f'Fold {i}', level=2)
            cm = ds_cv / f'fold_{i}_svm_confusion.png'
            if cm.exists():
                doc.add_paragraph('混淆矩阵（SVM）：行=真实标签，列=预测标签。数字为样本数。')
                doc.add_picture(str(cm), width=Inches(5))
            for model in ['svm','nb']:
                pr = ds_cv / f'fold_{i}_{model}_pr.png'
                roc = ds_cv / f'fold_{i}_{model}_roc.png'
                if pr.exists():
                    doc.add_paragraph(f'{model.upper()} Precision-Recall 曲线（此图用于观察模型在不同阈值下的精/召表现）')
                    doc.add_picture(str(pr), width=Inches(5))
                if roc.exists():
                    doc.add_paragraph(f'{model.upper()} ROC 曲线（用于衡量二分类器整体判别能力，AUC 越接近 1 越好）')
                    doc.add_picture(str(roc), width=Inches(5))
    else:
        doc.add_paragraph('No CV output folder for ' + ds)

# Samples for annotation section
doc.add_page_break()
doc.add_heading('样本抽样（供人工标注）', level=1)
for samp in ['samples_for_annotation_hotel.csv','samples_for_annotation_ecommerce.csv','samples_for_annotation_waimai.csv']:
    p = OUT / samp
    if p.exists():
        doc.add_heading(samp, level=2)
        # show first 10 lines as examples
        with p.open('r', encoding='utf-8') as f:
            for i,line in enumerate(f):
                if i >= 10:
                    break
                doc.add_paragraph(line.strip())

# Appendix: list files
doc.add_page_break()
doc.add_heading('附录：重要输出文件列表', level=1)
# list of important files
files = [
    'data/processed_hotel.csv', 'data/processed_ecommerce.csv', 'data/processed_waimai.csv',
    'output/labels_hotel.csv','output/labels_ecommerce.csv','output/labels_waimai.csv',
    'output/classification_report_hotel_nb.txt','output/classification_report_hotel_svm.txt',
    'output/classification_report_ecommerce_nb.txt','output/classification_report_ecommerce_svm.txt',
    'output/classification_report_waimai_nb.txt','output/classification_report_waimai_svm.txt',
    'output/cv/* (cv images and metrics)',
    'output/samples_for_annotation_*.csv',
    'output/experiment_report.md', 'output/experiment_report.docx'
]
for f in files:
    doc.add_paragraph(f)

# Save final
doc.save(FINAL)
print('Wrote', FINAL)
