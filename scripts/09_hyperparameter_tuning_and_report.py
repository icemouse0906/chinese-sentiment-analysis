import os
import pandas as pd
import numpy as np
from sklearn.model_selection import GridSearchCV, StratifiedKFold
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.naive_bayes import MultinomialNB
from sklearn.svm import LinearSVC
from sklearn.metrics import classification_report, roc_curve, precision_recall_curve, confusion_matrix, auc
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches
import warnings
warnings.filterwarnings('ignore')

# 数据集路径
DATASETS = {
    '电商': 'NLP数据集/电商评论数据/online_shopping_10_cats.csv',
    '酒店': 'NLP数据集/酒店评论数据/ChnSentiCorp_htl_all.csv',
    '外卖': 'NLP数据集/外卖评论数据/waimai_10k.csv'
}

# 模型与参数网格
MODELS = {
    'SVM': (LinearSVC(), {'C': [0.1, 1, 5, 10], 'max_iter': [1000, 3000]}),
    'NB': (MultinomialNB(), {'alpha': [0.1, 0.5, 1.0, 2.0]})
}

# 文档路径
REPORT_PATH = 'experiment_report_final.docx'

# 读取数据

def load_data(path):
    # 自动编码检测与回退
    encodings = ['utf-8', 'gbk', 'gb18030']
    for enc in encodings:
        try:
            df = pd.read_csv(path, encoding=enc)
            break
        except UnicodeDecodeError:
            continue
    else:
        raise RuntimeError(f'无法读取 {path}，请检查文件编码')
    # 自动检测标签列
    label_col = [c for c in df.columns if 'label' in c or '类别' in c or 'sentiment' in c]
    text_col = [c for c in df.columns if 'text' in c or '内容' in c or 'review' in c or '评论' in c]
    label = label_col[0] if label_col else df.columns[-1]
    text = text_col[0] if text_col else df.columns[0]
    # 删除标签为NaN的样本
    df = df.dropna(subset=[label])
    return df[text].astype(str), df[label]

# 训练与调优

def tune_and_evaluate(X, y, model, param_grid, dataset_name, model_name):
    # 检查每个类别样本数
    value_counts = pd.Series(y).value_counts()
    if (value_counts < 5).any():
        # 样本数不足，跳过
        msg = f"{dataset_name}-{model_name} 部分类别样本数不足5，无法进行5折交叉验证，已跳过。"
        print(msg)
        # 构造空报告
        report_df = pd.DataFrame({'info': [msg]})
        return None, report_df
    # 向量化
    tfidf = TfidfVectorizer(max_features=5000)
    X_vec = tfidf.fit_transform(X)
    # 交叉验证
    skf = StratifiedKFold(n_splits=5, shuffle=True, random_state=42)
    grid = GridSearchCV(model, param_grid, cv=skf, scoring='f1_weighted', n_jobs=-1)
    grid.fit(X_vec, y)
    best_model = grid.best_estimator_
    best_params = grid.best_params_
    # 评估
    y_pred = best_model.predict(X_vec)
    report = classification_report(y, y_pred, output_dict=True)
    # 保存分类报告
    report_df = pd.DataFrame(report).transpose()
    report_path = f'results/{dataset_name}_{model_name}_report.csv'
    os.makedirs('results', exist_ok=True)
    report_df.to_csv(report_path, encoding='utf-8')
    # 混淆矩阵
    cm = confusion_matrix(y, y_pred)
    plt.figure(figsize=(5,4))
    plt.imshow(cm, cmap='Blues')
    plt.title(f'{dataset_name}-{model_name} 混淆矩阵')
    plt.xlabel('预测')
    plt.ylabel('真实')
    plt.colorbar()
    plt.savefig(f'results/{dataset_name}_{model_name}_cm.png', bbox_inches='tight')
    plt.close()
    # PR/ROC 曲线（仅二分类）
    if len(set(y)) == 2:
        y_score = best_model.decision_function(X_vec) if hasattr(best_model, 'decision_function') else best_model.predict_proba(X_vec)[:,1]
        fpr, tpr, _ = roc_curve(y, y_score)
        roc_auc = auc(fpr, tpr)
        plt.figure()
        plt.plot(fpr, tpr, label=f'ROC curve (AUC={roc_auc:.2f})')
        plt.xlabel('False Positive Rate')
        plt.ylabel('True Positive Rate')
        plt.title(f'{dataset_name}-{model_name} ROC曲线')
        plt.legend()
        plt.savefig(f'results/{dataset_name}_{model_name}_roc.png', bbox_inches='tight')
        plt.close()
        precision, recall, _ = precision_recall_curve(y, y_score)
        pr_auc = auc(recall, precision)
        plt.figure()
        plt.plot(recall, precision, label=f'PR curve (AUC={pr_auc:.2f})')
        plt.xlabel('Recall')
        plt.ylabel('Precision')
        plt.title(f'{dataset_name}-{model_name} PR曲线')
        plt.legend()
        plt.savefig(f'results/{dataset_name}_{model_name}_pr.png', bbox_inches='tight')
        plt.close()
    # 返回结果
    return best_params, report_df

# 追加到Word报告

def append_to_report(results):
    if not os.path.exists(REPORT_PATH):
        doc = Document()
    else:
        doc = Document(REPORT_PATH)
    doc.add_heading('模型优化与性能对比', level=1)
    for dataset_name, model_name, best_params, report_df in results:
        doc.add_heading(f'{dataset_name} - {model_name} 超参数调优结果', level=2)
        if best_params is None:
            # 样本数不足说明
            doc.add_paragraph(report_df.iloc[0,0])
            continue
        doc.add_paragraph(f'最佳参数: {best_params}')
        # 分类报告表格
        table = doc.add_table(rows=1, cols=len(report_df.columns)+1)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = '类别'
        for i, col in enumerate(report_df.columns):
            hdr_cells[i+1].text = col
        for idx, row in report_df.iterrows():
            cells = table.add_row().cells
            cells[0].text = str(idx)
            for i, val in enumerate(row):
                cells[i+1].text = f'{val:.3f}' if isinstance(val, float) else str(val)
        # 插入混淆矩阵、PR/ROC图
        img_cm = f'results/{dataset_name}_{model_name}_cm.png'
        if os.path.exists(img_cm):
            doc.add_picture(img_cm, width=Inches(3.5))
        img_roc = f'results/{dataset_name}_{model_name}_roc.png'
        if os.path.exists(img_roc):
            doc.add_picture(img_roc, width=Inches(3.5))
        img_pr = f'results/{dataset_name}_{model_name}_pr.png'
        if os.path.exists(img_pr):
            doc.add_picture(img_pr, width=Inches(3.5))
    doc.save(REPORT_PATH)

# 主流程

def main():
    results = []
    for dataset_name, path in DATASETS.items():
        if not os.path.exists(path):
            print(f'数据集 {dataset_name} 未找到，跳过')
            continue
        X, y = load_data(path)
        for model_name, (model, param_grid) in MODELS.items():
            print(f'正在调优: {dataset_name} - {model_name}')
            best_params, report_df = tune_and_evaluate(X, y, model, param_grid, dataset_name, model_name)
            results.append((dataset_name, model_name, best_params, report_df))
    append_to_report(results)
    print('超参数调优与报告追加完成！')

if __name__ == '__main__':
    main()
